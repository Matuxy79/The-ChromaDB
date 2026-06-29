#!/usr/bin/env python3
"""Generate the legacy synthetic multi-pack test corpus for the CLS RAG+CAG stack.

This generator predates the current beamline-scoped corpus layout. It builds six
large synthetic topic packs, ~100 MB each (5 files per pack), deliberately spread
across every reader the pipeline supports (.pdf .docx .txt .md .html .csv .tsv
.json) so regression testing can still exercise the ingestion pipeline end to end.

Content is synthetic but topic-coherent: each pack draws on its own term bank and
sentence templates, so metadata-filtered retrieval still returns sensible,
on-topic evidence. Generation is seeded and deterministic.

Usage
-----
    python scripts/generate_test_corpus.py                 # full ~600 MB build
    python scripts/generate_test_corpus.py --scale 0.05    # ~30 MB smoke build
    python scripts/generate_test_corpus.py --only chemistry physics
    python scripts/generate_test_corpus.py --out data/corpus

Each synthetic pack lands in <out>/<pack_slug>/. Those slugs are retained for
legacy regression runs and can still be indexed by scripts/ingest_corpus.py,
but they do not describe the current CLS beamline taxonomy.
"""

from __future__ import annotations

import argparse
import json
import random
import sys
from pathlib import Path

MB = 1024 * 1024
DEFAULT_OUT = Path(__file__).resolve().parent.parent / "data" / "corpus"

# --------------------------------------------------------------------------- #
# Legacy synthetic content packs
# --------------------------------------------------------------------------- #
# Each pack supplies a term bank used to fill sentence templates (prose formats)
# and a tabular schema (columns + a row factory) used for csv/tsv/json.

PACKS: dict[str, dict] = {
    "chemistry": {
        "title": "Chemistry",
        "topics": [
            "reaction kinetics", "thermodynamics", "organic synthesis", "catalysis",
            "spectroscopy", "electrochemistry", "coordination chemistry",
            "molecular orbital theory", "acid-base equilibria", "polymer chemistry",
        ],
        "terms": [
            "activation energy", "enthalpy", "entropy", "Gibbs free energy",
            "rate constant", "equilibrium constant", "oxidation state", "ligand",
            "enantiomer", "stoichiometry", "molar mass", "electronegativity",
            "transition state", "nucleophile", "electrophile", "buffer capacity",
            "redox potential", "crystal field splitting", "resonance structure",
        ],
        "methods": [
            "nuclear magnetic resonance", "mass spectrometry", "X-ray diffraction",
            "gas chromatography", "titration", "calorimetry", "cyclic voltammetry",
        ],
        "units": ["kJ/mol", "mol/L", "K", "pH units", "g/mol", "nm", "ppm"],
        "table": {
            "columns": ["compound", "formula", "molar_mass_g_mol", "melting_point_C",
                        "solubility_g_L", "hazard_class"],
            "row": lambda r, i: [
                f"{r.choice(['sodium','potassium','calcium','ferric','cupric','zinc'])}-"
                f"{r.choice(['chloride','sulfate','nitrate','carbonate','phosphate'])}-{i}",
                f"{r.choice(['Na','K','Ca','Fe','Cu','Zn'])}{r.randint(1,3)}"
                f"{r.choice(['Cl','SO4','NO3','CO3','PO4'])}{r.randint(1,3)}",
                round(r.uniform(40, 480), 2),
                round(r.uniform(-90, 1200), 1),
                round(r.uniform(0.01, 900), 3),
                r.choice(["corrosive", "oxidizer", "irritant", "flammable", "inert"]),
            ],
        },
    },
    "computer_science": {
        "title": "Computer Science",
        "topics": [
            "distributed systems", "machine learning", "compilers", "computer networks",
            "operating systems", "cryptography", "database systems", "computational complexity",
            "graph algorithms", "concurrency", "information retrieval", "type theory",
        ],
        "terms": [
            "asymptotic complexity", "amortized cost", "hash table", "B-tree",
            "consensus protocol", "gradient descent", "backpropagation", "regularization",
            "context-free grammar", "register allocation", "cache coherence", "deadlock",
            "public-key exchange", "vector embedding", "inverted index", "lambda calculus",
            "garbage collection", "load balancing", "eventual consistency",
        ],
        "methods": [
            "dynamic programming", "divide and conquer", "memoization", "MapReduce",
            "stochastic optimization", "Monte Carlo sampling", "static analysis",
        ],
        "units": ["ms", "ops/s", "MB", "bits", "requests/s", "GFLOPs"],
        "table": {
            "columns": ["algorithm", "language", "time_complexity", "space_complexity",
                        "benchmark_ms", "dataset_size"],
            "row": lambda r, i: [
                f"{r.choice(['quick','merge','heap','radix','bucket','tim'])}sort_v{i}",
                r.choice(["Python", "Rust", "C++", "Go", "Java", "Haskell"]),
                r.choice(["O(n log n)", "O(n)", "O(n^2)", "O(log n)", "O(n*m)"]),
                r.choice(["O(1)", "O(n)", "O(log n)", "O(n^2)"]),
                round(r.uniform(0.4, 9800), 2),
                r.choice([1000, 10_000, 100_000, 1_000_000, 10_000_000]),
            ],
        },
    },
    "biology": {
        "title": "Biology",
        "topics": [
            "molecular biology", "genetics", "cell signalling", "evolutionary biology",
            "neuroscience", "immunology", "ecology", "microbiology", "developmental biology",
            "structural biology", "systems biology", "physiology",
        ],
        "terms": [
            "transcription factor", "messenger RNA", "allele frequency", "natural selection",
            "action potential", "synaptic plasticity", "antigen presentation", "homeostasis",
            "mitochondrial membrane", "gene expression", "protein folding", "enzyme kinetics",
            "phylogenetic tree", "apoptosis", "cytokine signalling", "DNA methylation",
            "ribosomal assembly", "trophic cascade", "osmoregulation",
        ],
        "methods": [
            "polymerase chain reaction", "CRISPR-Cas9 editing", "flow cytometry",
            "RNA sequencing", "cryo-electron microscopy", "patch-clamp recording", "Western blot",
        ],
        "units": ["mM", "µm", "mV", "copies/cell", "kDa", "CFU/mL", "bp"],
        "table": {
            "columns": ["gene", "organism", "chromosome", "function", "expression_level",
                        "conservation_score"],
            "row": lambda r, i: [
                f"{r.choice(['BRCA','TP','EGFR','MYC','SOD','HOX'])}{r.randint(1,9)}-{i}",
                r.choice(["H. sapiens", "M. musculus", "D. melanogaster",
                          "S. cerevisiae", "A. thaliana", "C. elegans"]),
                r.randint(1, 23),
                r.choice(["DNA repair", "cell cycle", "apoptosis", "metabolism",
                          "signal transduction", "transcription"]),
                round(r.uniform(0.1, 5000), 2),
                round(r.uniform(0.0, 1.0), 3),
            ],
        },
    },
    "physics": {
        "title": "Physics",
        "topics": [
            "quantum mechanics", "general relativity", "statistical mechanics",
            "condensed matter", "particle physics", "electromagnetism", "optics",
            "fluid dynamics", "astrophysics", "thermodynamics", "nonlinear dynamics",
            "plasma physics",
        ],
        "terms": [
            "wavefunction", "eigenstate", "gauge symmetry", "spacetime curvature",
            "partition function", "phase transition", "Fermi surface", "cross section",
            "Hamiltonian", "angular momentum", "renormalization", "entanglement entropy",
            "dispersion relation", "boundary condition", "perturbation series",
            "Lorentz invariance", "blackbody spectrum", "Bose-Einstein condensate",
        ],
        "methods": [
            "perturbation theory", "lattice simulation", "spectroscopic analysis",
            "interferometry", "Monte Carlo integration", "tensor network contraction",
            "synchrotron diffraction",
        ],
        "units": ["eV", "GeV", "K", "T", "Pa", "Hz", "m/s"],
        "table": {
            "columns": ["particle", "mass_MeV", "charge_e", "spin", "lifetime_s",
                        "interaction"],
            "row": lambda r, i: [
                f"{r.choice(['pion','kaon','muon','tau','quark','boson'])}-{i}",
                round(r.uniform(0.5, 175_000), 3),
                r.choice([-1, 0, 1, 2, -2]),
                r.choice(["0", "1/2", "1", "3/2", "2"]),
                f"{r.uniform(1e-25, 1e-6):.3e}",
                r.choice(["strong", "weak", "electromagnetic", "gravitational"]),
            ],
        },
    },
    "mathematics": {
        "title": "Mathematics",
        "topics": [
            "real analysis", "abstract algebra", "topology", "number theory",
            "differential geometry", "probability theory", "combinatorics",
            "functional analysis", "category theory", "graph theory", "measure theory",
            "algebraic geometry",
        ],
        "terms": [
            "compact set", "Hausdorff space", "ring homomorphism", "eigenvalue",
            "measure-preserving map", "convergence criterion", "Cauchy sequence",
            "isomorphism", "manifold", "group action", "prime ideal", "Banach space",
            "Lebesgue integral", "generating function", "spectral decomposition",
            "fixed point", "boundary operator", "uniform continuity", "convex hull",
        ],
        "methods": [
            "proof by induction", "contradiction argument", "diagonalization",
            "fixed-point iteration", "spectral decomposition", "epsilon-delta argument",
            "the probabilistic method",
        ],
        "units": ["", "dimensions", "iterations", "terms", "degrees of freedom"],
        "table": {
            "columns": ["theorem", "field", "year_proved", "difficulty", "open_problems",
                        "citations"],
            "row": lambda r, i: [
                f"Lemma-{i}-{r.choice(['fixed','dense','compact','prime','convex'])}",
                r.choice(["analysis", "algebra", "topology", "number theory",
                          "geometry", "probability"]),
                r.randint(1850, 2025),
                r.choice(["undergraduate", "graduate", "research", "open"]),
                r.randint(0, 12),
                r.randint(0, 9000),
            ],
        },
    },
    "literature": {
        "title": "Literature",
        "topics": [
            "the Victorian novel", "modernist poetry", "tragic drama", "the gothic tradition",
            "narrative theory", "postcolonial fiction", "the bildungsroman", "epic verse",
            "literary realism", "stream of consciousness", "satire", "Romanticism",
        ],
        "terms": [
            "unreliable narrator", "dramatic irony", "free indirect discourse",
            "extended metaphor", "iambic pentameter", "the picaresque", "epistolary form",
            "pathetic fallacy", "in medias res", "the frame narrative", "motif",
            "allegory", "interior monologue", "the tragic flaw", "bildungsroman",
            "verse drama", "imagery", "foreshadowing", "the omniscient narrator",
        ],
        "methods": [
            "close reading", "comparative analysis", "historicist reading",
            "structuralist analysis", "reader-response criticism", "narratological mapping",
            "archetypal criticism",
        ],
        "units": ["lines", "stanzas", "chapters", "acts", "cantos"],
        "table": {
            "columns": ["work", "author", "year", "genre", "narrative_mode", "themes"],
            "row": lambda r, i: [
                f"The {r.choice(['Silent','Distant','Hollow','Gilded','Autumn'])} "
                f"{r.choice(['House','Road','Letter','River','Garden'])} #{i}",
                f"{r.choice(['E.','C.','M.','J.','A.'])} "
                f"{r.choice(['Hartley','Vance','Crowe','Ashby','Quill'])}",
                r.randint(1798, 1965),
                r.choice(["novel", "poetry", "drama", "essay", "short story"]),
                r.choice(["first person", "omniscient", "epistolary",
                          "stream of consciousness", "dramatic"]),
                r.choice(["loss", "memory", "ambition", "exile", "redemption", "decay"]),
            ],
        },
    },
}

# Per-pack file plan: (stem, format, target_megabytes). Sizes sum to ~100 MB
# and are arranged so every supported reader appears across the corpus, while
# pdf/docx stay paper-sized (they are the slow, format-sensitive ones).
FILE_PLANS: dict[str, list[tuple[str, str, float]]] = {
    "chemistry": [
        ("review_reaction_kinetics", "md", 52),
        ("compound_property_table", "csv", 35),
        ("reaction_dataset", "json", 9),
        ("synthesis_protocol", "pdf", 3),
        ("lab_notebook", "docx", 1),
    ],
    "computer_science": [
        ("survey_distributed_systems", "html", 50),
        ("benchmark_results", "tsv", 35),
        ("algorithm_catalog", "json", 12),
        ("design_notes", "txt", 2),
        ("readme_pipeline", "md", 1),
    ],
    "biology": [
        ("lecture_notes_molecular_biology", "txt", 52),
        ("gene_expression_table", "csv", 33),
        ("assay_measurements", "tsv", 11),
        ("methods_paper", "pdf", 3),
        ("supplementary_figures", "html", 1),
    ],
    "physics": [
        ("monograph_quantum_mechanics", "md", 52),
        ("particle_catalog", "json", 33),
        ("experiment_log", "txt", 11),
        ("beamline_paper", "pdf", 3),
        ("analysis_memo", "docx", 1),
    ],
    "mathematics": [
        ("theorem_index", "tsv", 52),
        ("course_notes_analysis", "html", 34),
        ("conjecture_table", "csv", 11),
        ("proof_companion", "md", 2),
        ("seminar_minutes", "docx", 1),
    ],
    "literature": [
        ("critical_essays_collected", "txt", 52),
        ("anthology_commentary", "md", 33),
        ("study_guide", "html", 11),
        ("monograph_the_novel", "pdf", 3),
        ("reading_group_notes", "docx", 1),
    ],
}


# --------------------------------------------------------------------------- #
# Prose generation
# --------------------------------------------------------------------------- #
SENTENCE_TEMPLATES = [
    "The study of {topic} depends critically on {term}, which constrains {term2} under realistic conditions.",
    "Using {method}, researchers quantified {term} at roughly {val} {unit} across {n} independent trials.",
    "A central result connects {term} to {term2}, clarifying why {topic} resists naive treatment.",
    "When {term} dominates, {method} reveals a measurable shift in {term2} on the order of {val} {unit}.",
    "Section {n} revisits {topic}: the interplay between {term} and {term2} is examined in detail.",
    "Contrary to early models, {term} does not fully determine {term2}; {method} exposes the residual structure.",
    "We model {topic} by treating {term} as the controlling variable and {term2} as a derived quantity.",
    "Empirically, {term} scales with {term2}, and {method} bounds the deviation below {val} {unit}.",
    "The {topic} literature debates whether {term} or {term2} better predicts the observed {val} {unit} response.",
    "Figure {n} summarises how {term} propagates through the system before {method} stabilises {term2}.",
    "Repeated application of {method} narrows the uncertainty in {term} to within {val} {unit}.",
    "An open question in {topic} is whether {term} and {term2} remain coupled outside the {val} {unit} regime.",
    "For pedagogical clarity, {term} is introduced first, then related to {term2} through {method}.",
    "The dataset records {n} cases in which {term} exceeded {term2}, each verified by {method}.",
    "Taken together, these observations recast {topic} as a balance between {term} and {term2}.",
]


def _fill(r: random.Random, pack: dict, template: str) -> str:
    return template.format(
        topic=r.choice(pack["topics"]),
        term=r.choice(pack["terms"]),
        term2=r.choice(pack["terms"]),
        method=r.choice(pack["methods"]),
        unit=r.choice(pack["units"]) or "units",
        val=round(r.uniform(0.01, 9999), 2),
        n=r.randint(1, 400),
    )


def paragraph(r: random.Random, pack: dict, sentences: int = 6) -> str:
    return " ".join(_fill(r, pack, r.choice(SENTENCE_TEMPLATES)) for _ in range(sentences))


def _estimate_unit_bytes(make_unit, samples: int = 64) -> float:
    r = random.Random(0)
    total = sum(len(make_unit(r, i).encode("utf-8")) for i in range(samples))
    return max(total / samples, 1.0)


def build_to_size(make_unit, target_bytes: int, joiner: str, r: random.Random) -> str:
    """Append units (paragraphs/rows/...) until the joined text reaches target_bytes."""
    avg = _estimate_unit_bytes(make_unit) + len(joiner.encode("utf-8"))
    estimate = max(int(target_bytes / avg), 1)
    parts = [make_unit(r, i) for i in range(estimate)]
    text = joiner.join(parts)
    i = estimate
    while len(text.encode("utf-8")) < target_bytes:
        block = [make_unit(r, j) for j in range(i, i + 256)]
        text += joiner + joiner.join(block)
        i += 256
    return text


# --------------------------------------------------------------------------- #
# Format writers
# --------------------------------------------------------------------------- #
def write_txt(path: Path, pack: dict, target: int, r: random.Random) -> None:
    text = build_to_size(lambda r, i: paragraph(r, pack), target, "\n\n", r)
    header = f"{pack['title']} — Working Notes and Lecture Material\n{'=' * 60}\n\n"
    path.write_text(header + text, encoding="utf-8")


def write_md(path: Path, pack: dict, target: int, r: random.Random) -> None:
    def unit(r: random.Random, i: int) -> str:
        if i % 12 == 0:
            return f"## {r.choice(pack['topics']).title()} — note {i}\n\n{paragraph(r, pack)}"
        if i % 7 == 0:
            items = "\n".join(f"- **{r.choice(pack['terms'])}**: {_fill(r, pack, SENTENCE_TEMPLATES[1])}"
                              for _ in range(4))
            return items
        return paragraph(r, pack)

    body = build_to_size(unit, target, "\n\n", r)
    path.write_text(f"# {pack['title']} — Reference Compendium\n\n{body}\n", encoding="utf-8")


def write_html(path: Path, pack: dict, target: int, r: random.Random) -> None:
    def unit(r: random.Random, i: int) -> str:
        if i % 10 == 0:
            return f"<h2>{r.choice(pack['topics']).title()} (section {i})</h2>"
        return f"<p>{paragraph(r, pack)}</p>"

    body = build_to_size(unit, target, "\n", r)
    html_doc = (
        "<!DOCTYPE html>\n<html lang=\"en\">\n<head><meta charset=\"utf-8\">"
        f"<title>{pack['title']} survey</title></head>\n<body>\n"
        f"<h1>{pack['title']} — Annotated Survey</h1>\n{body}\n</body>\n</html>\n"
    )
    path.write_text(html_doc, encoding="utf-8")


def write_csv(path: Path, pack: dict, target: int, r: random.Random, delim: str = ",") -> None:
    cols = pack["table"]["columns"]
    row_fn = pack["table"]["row"]
    header = delim.join(cols)

    def unit(r: random.Random, i: int) -> str:
        return delim.join(str(v) for v in row_fn(r, i))

    body = build_to_size(unit, max(target - len(header), 1), "\n", r)
    path.write_text(f"{header}\n{body}\n", encoding="utf-8")


def write_tsv(path: Path, pack: dict, target: int, r: random.Random) -> None:
    write_csv(path, pack, target, r, delim="\t")


def write_json(path: Path, pack: dict, target: int, r: random.Random) -> None:
    cols = pack["table"]["columns"]
    row_fn = pack["table"]["row"]
    # Estimate records from one serialized sample, then dump as a single array.
    sample = dict(zip(cols, row_fn(random.Random(0), 0)))
    per = len(json.dumps(sample).encode("utf-8")) + 2
    count = max(int(target / per), 1)
    records = [dict(zip(cols, row_fn(r, i))) for i in range(count)]
    payload = {
        "discipline": pack["title"],
        "schema": cols,
        "record_count": len(records),
        "records": records,
    }
    path.write_text(json.dumps(payload, indent=1), encoding="utf-8")


def write_pdf(path: Path, pack: dict, target: int, r: random.Random) -> None:
    import fitz  # PyMuPDF

    def make_page_text(i: int) -> str:
        head = f"{pack['title']} — Section {i}: {r.choice(pack['topics']).title()}\n\n"
        return head + "\n\n".join(paragraph(r, pack, sentences=5) for _ in range(7))

    def render(n_pages: int) -> bytes:
        doc = fitz.open()
        rect = fitz.Rect(56, 56, 539, 785)  # A4 text frame
        for i in range(n_pages):
            page = doc.new_page(width=595, height=842)
            page.insert_textbox(rect, make_page_text(i), fontsize=9, fontname="helv")
        data = doc.tobytes()
        doc.close()
        return data

    # Calibrate bytes/page on a small sample, then render to target.
    calib_pages = 24
    sample = render(calib_pages)
    per_page = max(len(sample) / calib_pages, 1.0)
    n_pages = max(int(target / per_page), 1)
    path.write_bytes(render(n_pages))


def write_docx(path: Path, pack: dict, target: int, r: random.Random) -> None:
    import docx  # python-docx

    def render(n_paras: int) -> bytes:
        document = docx.Document()
        document.add_heading(f"{pack['title']} — Working Document", level=0)
        for i in range(n_paras):
            if i % 15 == 0:
                document.add_heading(f"{r.choice(pack['topics']).title()} (note {i})", level=1)
            document.add_paragraph(paragraph(r, pack, sentences=6))
        import io
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    # Calibrate compressed bytes/paragraph, then render to target (capped for speed).
    calib = 200
    sample = render(calib)
    per_para = max(len(sample) / calib, 1.0)
    n_paras = min(max(int(target / per_para), 1), 40_000)
    path.write_bytes(render(n_paras))


WRITERS = {
    "txt": write_txt,
    "md": write_md,
    "html": write_html,
    "csv": write_csv,
    "tsv": write_tsv,
    "json": write_json,
    "pdf": write_pdf,
    "docx": write_docx,
}


# --------------------------------------------------------------------------- #
# Driver
# --------------------------------------------------------------------------- #
def human(n: int) -> str:
    f = float(n)
    for unit in ("B", "KB", "MB", "GB"):
        if f < 1024 or unit == "GB":
            return f"{f:.1f} {unit}"
        f /= 1024
    return f"{f:.1f} GB"


def generate(out: Path, only: list[str] | None, scale: float, seed: int) -> None:
    disciplines = only or list(FILE_PLANS)
    grand_total = 0
    for discipline in disciplines:
        if discipline not in PACKS:
            print(f"  ! unknown legacy pack {discipline!r} — skipping", file=sys.stderr)
            continue
        pack = PACKS[discipline]
        folder = out / discipline
        folder.mkdir(parents=True, exist_ok=True)
        print(f"\n[{discipline}]")
        disc_total = 0
        for idx, (stem, fmt, mb) in enumerate(FILE_PLANS[discipline]):
            target = int(mb * MB * scale)
            r = random.Random(f"{seed}:{discipline}:{stem}")
            dest = folder / f"{stem}.{fmt}"
            WRITERS[fmt](dest, pack, target, r)
            size = dest.stat().st_size
            disc_total += size
            print(f"  {dest.name:<38} {fmt:<5} {human(size):>10}")
        print(f"  {'— pack total —':<38} {'':5} {human(disc_total):>10}")
        grand_total += disc_total
    print(f"\nTOTAL written: {human(grand_total)} under {out}")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT,
                        help="Output corpus root (default: data/corpus).")
    parser.add_argument("--only", nargs="*", default=None,
                        help="Limit to these legacy synthetic packs (default: all six).")
    parser.add_argument("--scale", type=float, default=1.0,
                        help="Multiply every target size (e.g. 0.05 for a quick smoke build).")
    parser.add_argument("--seed", type=int, default=1729, help="RNG seed for reproducibility.")
    args = parser.parse_args()
    generate(args.out, args.only, args.scale, args.seed)


if __name__ == "__main__":
    main()
