"""Document readers for the iron triangle: small, local, no huge dependencies.

Each reader returns a list of (page_number, text) tuples so the rest of the
pipeline (chunking, embedding, indexing) stays document-type-agnostic.
"""

from __future__ import annotations

import csv
import html
import json
import re
from pathlib import Path
from typing import Any

import fitz


def _norm_whitespace(text: str) -> str:
    """Collapse whitespace but keep paragraph breaks."""
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


_GUTENBERG_START = re.compile(
    r"\*\*\*\s*START OF (?:THE|THIS) PROJECT GUTENBERG EBOOK.*?\*\*\*",
    re.IGNORECASE | re.DOTALL,
)
_GUTENBERG_END = re.compile(
    r"\*\*\*\s*END OF (?:THE|THIS) PROJECT GUTENBERG EBOOK.*?\*\*\*",
    re.IGNORECASE | re.DOTALL,
)


def _strip_gutenberg_boilerplate(text: str) -> str:
    """Drop Project Gutenberg license/header text when present."""
    start = _GUTENBERG_START.search(text)
    if start:
        text = text[start.end():]
    end = _GUTENBERG_END.search(text)
    if end:
        text = text[:end.start()]
    return text


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def load_pdf(path: Path) -> list[tuple[int, str]]:
    pages: list[tuple[int, str]] = []
    with fitz.open(path) as document:
        for page_index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            if text:
                pages.append((page_index, _norm_whitespace(text)))
    return pages


# ---------------------------------------------------------------------------
# Plain text / Markdown
# ---------------------------------------------------------------------------
def load_text(path: Path) -> list[tuple[int, str]]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    text = _strip_gutenberg_boilerplate(text)
    return [(1, _norm_whitespace(text))] if text else []


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def load_docx(path: Path) -> list[tuple[int, str]]:
    try:
        import docx
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for .docx support. Install: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    # Tables often carry the real data in facility docs.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    if not paragraphs:
        return []

    # DOCX has no true pages; treat each top-level heading block as a pseudo-page
    # so section boundaries survive chunking.
    pages: list[tuple[int, str]] = []
    buffer: list[str] = []
    page_no = 1
    heading_re = re.compile(r"^(\d+(\.\d+)*\s+)?[A-Z][A-Za-z0-9\s\-]{2,}$")

    def flush() -> None:
        nonlocal buffer, page_no
        if buffer:
            pages.append((page_no, _norm_whitespace("\n\n".join(buffer))))
            page_no += 1
            buffer = []

    for para in paragraphs:
        if heading_re.match(para) and buffer:
            flush()
        buffer.append(para)
    flush()
    return pages


# ---------------------------------------------------------------------------
# HTML / HTM
# ---------------------------------------------------------------------------
def load_html(path: Path) -> list[tuple[int, str]]:
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self._chunks: list[str] = []
            self._skip = 0
            self._current: list[str] = []

        def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
            tag = tag.lower()
            if tag in {"script", "style", "nav", "footer", "header"}:
                self._skip += 1
            elif tag in {"p", "br", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6", "tr"}:
                self._flush_current()

        def handle_endtag(self, tag: str) -> None:
            tag = tag.lower()
            if tag in {"script", "style", "nav", "footer", "header"}:
                self._skip = max(0, self._skip - 1)
            elif tag in {"p", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6", "tr"}:
                self._flush_current()

        def handle_data(self, data: str) -> None:
            if self._skip:
                return
            self._current.append(data)

        def _flush_current(self) -> None:
            text = " ".join(self._current).strip()
            self._current = []
            if text:
                self._chunks.append(text)

        def text(self) -> str:
            self._flush_current()
            return "\n\n".join(self._chunks)

    raw = path.read_text(encoding="utf-8", errors="replace")
    extractor = _TextExtractor()
    try:
        extractor.feed(raw)
    except Exception:
        # Fallback: strip tags with a regex if the parser chokes on malformed HTML.
        text = re.sub(r"<script[^>]*>.*?</script>", " ", raw, flags=re.S | re.I)
        text = re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=re.S | re.I)
        text = re.sub(r"<[^>]+>", " ", text)
        text = html.unescape(text)
        return [(1, _norm_whitespace(text))]

    text = extractor.text()
    return [(1, _norm_whitespace(text))] if text else []


# ---------------------------------------------------------------------------
# CSV / TSV
# ---------------------------------------------------------------------------
def _load_delimited(path: Path, delimiter: str) -> list[tuple[int, str]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    if not text.strip():
        return []

    rows = list(csv.reader(text.splitlines(), delimiter=delimiter))
    if not rows:
        return []

    header = " | ".join(rows[0])
    pages: list[tuple[int, str]] = []
    for idx, row in enumerate(rows[1:], start=1):
        if not row:
            continue
        row_text = " | ".join(row)
        pages.append((idx, f"Header: {header}\nRow: {row_text}"))
    return pages


def load_csv(path: Path) -> list[tuple[int, str]]:
    return _load_delimited(path, ",")


def load_tsv(path: Path) -> list[tuple[int, str]]:
    return _load_delimited(path, "\t")


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------
def _flatten(obj: Any, prefix: str = "") -> list[str]:
    lines: list[str] = []
    if isinstance(obj, dict):
        for key, value in obj.items():
            new_prefix = f"{prefix}.{key}" if prefix else str(key)
            if isinstance(value, (dict, list)):
                lines.extend(_flatten(value, new_prefix))
            else:
                lines.append(f"{new_prefix}: {value}")
    elif isinstance(obj, list):
        for index, item in enumerate(obj):
            new_prefix = f"{prefix}[{index}]"
            if isinstance(item, (dict, list)):
                lines.extend(_flatten(item, new_prefix))
            else:
                lines.append(f"{new_prefix}: {item}")
    else:
        lines.append(f"{prefix}: {obj}")
    return lines


def load_json(path: Path) -> list[tuple[int, str]]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    if not raw.strip():
        return []
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON in {path.name}: {exc}") from exc

    lines = _flatten(data)
    if not lines:
        return []

    # Group flattened lines into pseudo-pages so chunking has boundaries.
    page_size = 20
    pages: list[tuple[int, str]] = []
    for i in range(0, len(lines), page_size):
        chunk = lines[i : i + page_size]
        pages.append((i // page_size + 1, "\n".join(chunk)))
    return pages


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------
SUPPORTED_SUFFIXES: set[str] = {
    ".pdf",
    ".txt",
    ".md",
    ".docx",
    ".html",
    ".htm",
    ".csv",
    ".tsv",
    ".json",
}

_READERS: dict[str, Any] = {
    ".pdf": load_pdf,
    ".txt": load_text,
    ".md": load_text,
    ".docx": load_docx,
    ".html": load_html,
    ".htm": load_html,
    ".csv": load_csv,
    ".tsv": load_tsv,
    ".json": load_json,
}


def load_document(path: Path) -> list[tuple[int, str]]:
    suffix = path.suffix.lower()
    reader = _READERS.get(suffix)
    if reader is None:
        raise ValueError(
            f"Unsupported file type: {suffix}. Supported: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
        )
    return reader(path)


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_SUFFIXES
