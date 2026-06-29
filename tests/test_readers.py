import json
import tempfile
import unittest
from pathlib import Path

from cls_backend.readers import (
    SUPPORTED_SUFFIXES,
    is_supported,
    load_csv,
    load_docx,
    load_html,
    load_json,
    load_text,
    load_tsv,
)


class TextReaderTests(unittest.TestCase):
    def test_load_text_returns_single_page(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        ) as handle:
            handle.write("Line one.\n\nLine two.\n")
            temp_path = Path(handle.name)

        try:
            pages = load_text(temp_path)
            self.assertEqual(len(pages), 1)
            self.assertEqual(pages[0][0], 1)
            self.assertIn("Line one", pages[0][1])
            self.assertIn("Line two", pages[0][1])
        finally:
            temp_path.unlink(missing_ok=True)

    def test_load_text_strips_gutenberg_boilerplate(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        ) as handle:
            handle.write(
                "Header contact info\n"
                "*** START OF THE PROJECT GUTENBERG EBOOK SAMPLE ***\n"
                "Actual story sentence.\n"
                "*** END OF THE PROJECT GUTENBERG EBOOK SAMPLE ***\n"
                "License contact links\n"
            )
            temp_path = Path(handle.name)

        try:
            pages = load_text(temp_path)
            self.assertEqual(len(pages), 1)
            text = pages[0][1]
            self.assertIn("Actual story sentence", text)
            self.assertNotIn("Header contact info", text)
            self.assertNotIn("License contact links", text)
        finally:
            temp_path.unlink(missing_ok=True)


class CsvReaderTests(unittest.TestCase):
    def test_load_csv_turns_rows_into_pages(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".csv", delete=False, encoding="utf-8"
        ) as handle:
            handle.write("name,role,extension\n")
            handle.write("Control Room,operations,3570\n")
            handle.write("Floor Coordinator,access,3639\n")
            temp_path = Path(handle.name)

        try:
            pages = load_csv(temp_path)
            self.assertEqual(len(pages), 2)
            self.assertEqual(pages[0][0], 1)
            self.assertIn("Header: name | role | extension", pages[0][1])
            self.assertIn("Control Room", pages[0][1])
            self.assertIn("3570", pages[0][1])
        finally:
            temp_path.unlink(missing_ok=True)


class TsvReaderTests(unittest.TestCase):
    def test_load_tsv_parses_tab_delimiter(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".tsv", delete=False, encoding="utf-8"
        ) as handle:
            handle.write("name\trole\textension\n")
            handle.write("Control Room\toperations\t3570\n")
            temp_path = Path(handle.name)

        try:
            pages = load_tsv(temp_path)
            self.assertEqual(len(pages), 1)
            self.assertIn("Control Room | operations | 3570", pages[0][1])
        finally:
            temp_path.unlink(missing_ok=True)


class JsonReaderTests(unittest.TestCase):
    def test_load_json_flattens_nested_data(self):
        payload = {
            "facility": "CLS",
            "contacts": {
                "control_room": {"extension": 3570},
                "floor_coordinator": {"extension": 3639},
            },
        }
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".json", delete=False, encoding="utf-8"
        ) as handle:
            json.dump(payload, handle)
            temp_path = Path(handle.name)

        try:
            pages = load_json(temp_path)
            self.assertGreaterEqual(len(pages), 1)
            combined = "\n".join(text for _, text in pages)
            self.assertIn("facility: CLS", combined)
            self.assertIn("contacts.control_room.extension: 3570", combined)
        finally:
            temp_path.unlink(missing_ok=True)

    def test_load_json_rejects_invalid_json(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".json", delete=False, encoding="utf-8"
        ) as handle:
            handle.write("not json")
            temp_path = Path(handle.name)

        try:
            with self.assertRaises(ValueError):
                load_json(temp_path)
        finally:
            temp_path.unlink(missing_ok=True)


class HtmlReaderTests(unittest.TestCase):
    def test_load_html_strips_tags_and_scripts(self):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".html", delete=False, encoding="utf-8"
        ) as handle:
            handle.write(
                "<html><head><script>alert('x')</script></head>"
                "<body><h1>CLS Contacts</h1>"
                "<p>Control Room: ext. 3570</p>"
                "<style>.x{color:red}</style>"
                "<footer>ignore me</footer></body></html>"
            )
            temp_path = Path(handle.name)

        try:
            pages = load_html(temp_path)
            self.assertEqual(len(pages), 1)
            text = pages[0][1]
            self.assertIn("CLS Contacts", text)
            self.assertIn("Control Room: ext. 3570", text)
            self.assertNotIn("<script>", text)
            self.assertNotIn("ignore me", text)
        finally:
            temp_path.unlink(missing_ok=True)


class DocxReaderTests(unittest.TestCase):
    def test_load_docx_extracts_paragraphs_and_tables(self):
        try:
            import docx
        except ImportError:
            self.skipTest("python-docx not installed")

        with tempfile.NamedTemporaryFile(
            suffix=".docx", delete=False
        ) as handle:
            temp_path = Path(handle.name)

        try:
            document = docx.Document()
            document.add_heading("Facility Contacts", level=1)
            document.add_paragraph("Control Room extension 3570.")
            document.add_paragraph("Floor coordinator extension 3639.")
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Role"
            table.rows[0].cells[1].text = "Extension"
            row = table.add_row().cells
            row[0].text = "Control Room"
            row[1].text = "3570"
            document.save(str(temp_path))

            pages = load_docx(temp_path)
            self.assertGreaterEqual(len(pages), 1)
            combined = "\n\n".join(text for _, text in pages)
            self.assertIn("Control Room extension 3570", combined)
            self.assertIn("Floor coordinator extension 3639", combined)
            self.assertIn("Role | Extension", combined)
            self.assertIn("Control Room | 3570", combined)
        finally:
            temp_path.unlink(missing_ok=True)


class SupportedSuffixTests(unittest.TestCase):
    def test_common_types_supported(self):
        self.assertIn(".pdf", SUPPORTED_SUFFIXES)
        self.assertIn(".txt", SUPPORTED_SUFFIXES)
        self.assertIn(".md", SUPPORTED_SUFFIXES)
        self.assertIn(".docx", SUPPORTED_SUFFIXES)
        self.assertIn(".html", SUPPORTED_SUFFIXES)
        self.assertIn(".csv", SUPPORTED_SUFFIXES)
        self.assertIn(".tsv", SUPPORTED_SUFFIXES)
        self.assertIn(".json", SUPPORTED_SUFFIXES)

    def test_is_supported_checks_case_insensitive_suffix(self):
        self.assertTrue(is_supported(Path("doc.PDF")))
        self.assertTrue(is_supported(Path("data.DOCX")))
        self.assertFalse(is_supported(Path("archive.zip")))


if __name__ == "__main__":
    unittest.main()
